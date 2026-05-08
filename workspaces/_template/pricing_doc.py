"""Generate Sequor Pricing Strategy Word document."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Calibri"

# Deck palette
D500 = RGBColor(0x29, 0x91, 0xAA)   # primary accent
D700 = RGBColor(0x17, 0x5A, 0x73)   # dark supporting
D800 = RGBColor(0x14, 0x47, 0x5D)   # darkest
BODY = RGBColor(0x4A, 0x5E, 0x6E)   # body text
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG = RGBColor(0xF4, 0xFB, 0xFD)
BORDER = RGBColor(0xAA, 0xE5, 0xE8)
MUTED = RGBColor(0x8A, 0x9E, 0xB0)

doc = Document()

# ── Page margins ────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Helper: set paragraph shading ────────────────────────────────────────────
def shade_paragraph(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            tag = OxmlElement(f'w:{edge}')
            tag.set(qn('w:val'), kwargs[edge].get('val', 'single'))
            tag.set(qn('w:sz'), str(kwargs[edge].get('sz', 4)))
            tag.set(qn('w:space'), '0')
            tag.set(qn('w:color'), kwargs[edge].get('color', 'auto'))
            tcBorders.append(tag)
    tcPr.append(tcBorders)


def add_run(para, text, bold=False, italic=False, size=11,
             color=BODY, font=FONT):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = font
    return r


# ── COVER ────────────────────────────────────────────────────────────────────
# Top accent bar
para = doc.add_paragraph()
para.paragraph_format.space_before = Pt(0)
para.paragraph_format.space_after = Pt(0)
shade_paragraph(para, '2991AA')
run = para.add_run('  ')
run.font.size = Pt(6)

# Title block
para = doc.add_paragraph()
para.paragraph_format.space_before = Pt(18)
para.paragraph_format.space_after = Pt(4)
add_run(para, 'Sequor', bold=True, size=36, color=D800)

p2 = doc.add_paragraph()
p2.paragraph_format.space_before = Pt(0)
p2.paragraph_format.space_after = Pt(2)
add_run(p2, 'Pricing Strategy', bold=True, size=22, color=D500)

p3 = doc.add_paragraph()
p3.paragraph_format.space_after = Pt(2)
add_run(p3, 'Investor Presentation — May 2026', italic=True, size=12, color=MUTED)

# Thin rule
hr = doc.add_paragraph()
hr.paragraph_format.space_before = Pt(4)
hr.paragraph_format.space_after = Pt(16)
pPr = hr._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bot = OxmlElement('w:bottom')
bot.set(qn('w:val'), 'single')
bot.set(qn('w:sz'), '6')
bot.set(qn('w:space'), '1')
bot.set(qn('w:color'), 'AAE5E8')
pBdr.append(bot)
pPr.append(pBdr)

# ── EXECUTIVE SUMMARY ─────────────────────────────────────────────────────────
h1 = doc.add_paragraph()
h1.paragraph_format.space_before = Pt(12)
h1.paragraph_format.space_after = Pt(6)
add_run(h1, 'Executive Summary', bold=True, size=16, color=D800)

exec_text = (
    "Sequor uses a four-tier freemium model designed to convert small teams quickly "
    "(Free → Solo) while capturing larger teams with compliance and analytics needs "
    "(Starter → Pro). Pricing is anchored to value delivered, not headcount alone: "
    "the Free tier creates geographic lock-in through the learning loop; the Starter tier "
    "is positioned as the default for professional services teams; the Pro tier targets "
    "larger teams with compliance obligations. All tiers include WhatsApp + Email, "
    "auto-reply AI, and the RAG Document Hub — differentiation lives in escalation "
    "controls, audit retention, and support depth."
)
pe = doc.add_paragraph()
pe.paragraph_format.space_after = Pt(14)
add_run(pe, exec_text, size=11, color=BODY)

# ── PRICING TABLE ────────────────────────────────────────────────────────────
h2 = doc.add_paragraph()
h2.paragraph_format.space_before = Pt(4)
h2.paragraph_format.space_after = Pt(8)
add_run(h2, 'Tier Overview', bold=True, size=16, color=D800)

tiers = [
    {
        'name': 'Free',
        'price': '$0',
        'period': 'month',
        'badge': None,
        'bg': 'F4FBFD',
        'header_color': '94A3B8',
        'accent': '64748B',
        'tagline': 'Get started — no commitment',
        'features': [
            '50 messages / month',
            'WhatsApp + Email',
            'Auto-reply AI',
            'RAG Document Hub',
            '1 operator',
        ],
        'why': (
            "Converts curious users into believers at zero cost. "
            "The 50-message ceiling is intentional — active teams hit it within days, "
            "motivating an upgrade without a hard sales conversation."
        ),
    },
    {
        'name': 'Solo',
        'price': '$15',
        'period': 'month',
        'badge': None,
        'bg': 'D4EAF4',
        'header_color': '7FBED4',
        'accent': '2991AA',
        'tagline': 'For independent operators',
        'features': [
            '200 messages / month',
            'Auto-reply + RAG',
            '30-day message history',
            '1 operator',
        ],
        'why': (
            "Targets the solo consultant or freelancer who represents the majority of "
            "SE Asian micro-businesses. 200 messages is enough for active practitioners "
            "while staying well below the threshold that makes Slack/WHATSAPP "
            "worthwhile. 30-day history prevents the 'where did that conversation go?' "
            "frustration that drives churn on free tiers."
        ),
    },
    {
        'name': 'Starter',
        'price': '$35',
        'period': 'month',
        'badge': 'Most Popular',
        'bg': 'D4EAF4',
        'header_color': '2991AA',
        'accent': '175A73',
        'tagline': 'Default choice for professional teams',
        'features': [
            'Unlimited messages*',
            'Smart Escalations',
            'PDPA audit trail',
            'Up to 3 operators',
            'Priority support',
        ],
        'why': (
            "The core revenue tier. Targets 3–5 person professional services teams "
            "(consultants, agencies, law firms, accounting practices) — the exact profile "
            "that suffers most from coverage gaps and can afford $35/seat. "
            "Smart Escalations justify the jump from Solo by solving a real pain: "
            "VIP contacts and refund threats that slip through without routing. "
            "PDPA audit trail is a near-mandatory requirement for Singapore-based firms. "
            "Unlimited messages removes a friction point that constantly nags at $15 tier users."
        ),
    },
    {
        'name': 'Pro',
        'price': '$55',
        'period': 'month',
        'badge': None,
        'bg': 'AAE5E8',
        'header_color': '175A73',
        'accent': '14475D',
        'tagline': 'For teams with compliance obligations',
        'features': [
            'Unlimited everything',
            'Advanced analytics',
            'Custom key phrase maps',
            'PDPA compliance report',
            'Dedicated support',
        ],
        'why': (
            "Targets larger professional services firms (10+ seats) or regulated industries "
            "where the PDPA compliance report and custom routing rules are non-negotiable "
            "requirements, not nice-to-haves. At $55/seat for a 5-person team, the $275/month "
            "ACV is credible for sales-assisted motion. Custom key phrase maps allow operators "
            "to encode firm-specific escalation rules — a direct productivity multiplier that "
            "justifies the price premium over Starter."
        ),
    },
]

col_widths = [Cm(3.5), Cm(6.0), Cm(5.5), Cm(4.5)]
table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

# Header row
hdr_cells = table.rows[0].cells
headers = ['Tier', 'Features', 'Rationale', 'Target Customer']
for i, (cell, hdr, w) in enumerate(zip(hdr_cells, headers, col_widths)):
    cell.width = w
    shade_cell(cell, '14475D')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(hdr)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = WHITE
    run.font.name = FONT

# Data rows
customers = [
    'Solo operators, freelancers',
    'Active 1-person businesses',
    '3–5 person professional services teams',
    'Larger or regulated firms (10+ seats)',
]
for idx, (tier, cust) in enumerate(zip(tiers, customers)):
    row = table.add_row()
    cells = row.cells
    bg = 'FFFFFF' if idx % 2 == 0 else 'F4FBFD'

    # Name + price
    name_cell = cells[0]
    name_cell.width = col_widths[0]
    shade_cell(name_cell, tier['bg'])
    p = name_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(tier['name'])
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor.from_string(tier['accent']); r.font.name = FONT
    p2 = name_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(tier['price'])
    r2.bold = True; r2.font.size = Pt(22); r2.font.color.rgb = RGBColor.from_string(tier['accent']); r2.font.name = FONT
    p3 = name_cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(f"/{tier['period']}")
    r3.font.size = Pt(10); r3.font.color.rgb = MUTED; r3.font.name = FONT
    if tier['badge']:
        p4 = name_cell.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r4 = p4.add_run(tier['badge'])
        r4.bold = True; r4.font.size = Pt(8); r4.font.color.rgb = WHITE; r4.font.name = FONT
        shade_paragraph(p4, '2991AA')

    # Features
    feat_cell = cells[1]
    feat_cell.width = col_widths[1]
    shade_cell(feat_cell, bg)
    for feat in tier['features']:
        fp = feat_cell.add_paragraph()
        fp.paragraph_format.space_before = Pt(1)
        fp.paragraph_format.space_after = Pt(1)
        fr = fp.add_run(f"• {feat}")
        fr.font.size = Pt(10); fr.font.color.rgb = BODY; fr.font.name = FONT

    # Rationale
    rat_cell = cells[2]
    rat_cell.width = col_widths[2]
    shade_cell(rat_cell, bg)
    rp = rat_cell.paragraphs[0]
    rr = rp.add_run(tier['why'])
    rr.font.size = Pt(10); rr.font.color.rgb = BODY; rr.font.name = FONT

    # Customer
    cust_cell = cells[3]
    cust_cell.width = col_widths[3]
    shade_cell(cust_cell, bg)
    cp = cust_cell.paragraphs[0]
    cr = cp.add_run(cust)
    cr.font.size = Pt(10); cr.font.color.rgb = BODY; cr.font.name = FONT

doc.add_paragraph()  # spacer

# ── PRICING PHILOSOPHY ────────────────────────────────────────────────────────
h3 = doc.add_paragraph()
h3.paragraph_format.space_before = Pt(10)
h3.paragraph_format.space_after = Pt(8)
add_run(h3, 'Pricing Philosophy', bold=True, size=16, color=D800)

philosophy_items = [
    ("Anchor to pain, not cost",
     "Teams pay to avoid the pain of missing a client message. The right price is "
     "whatever is less than the expected revenue loss from one missed escalation."),
    ("Freemium as acquisition, not revenue",
     "The Free tier is a customer acquisition channel, not a revenue line. "
     "Every free user is a future upsell candidate with demonstrated intent."),
    ("Per-seat keeps alignment",
     "Pricing per seat aligns Sequor's revenue with the customer's team growth. "
     "Larger teams pay more because they get more value."),
    ("Unlimited removes friction",
     "Once users are thinking about message counts, they're not thinking about "
     "their actual workflow. Unlimited on Starter+ removes a constant anxiety "
     "that drives churn on metered plans."),
    ("Compliance is a premium feature",
     "Teams with genuine PDPA obligations (legal, medical, financial) treat audit "
     "trails and compliance reports as infrastructure costs, not optional upgrades. "
     "Pro pricing reflects this."),
]

for title, body in philosophy_items:
    ph = doc.add_paragraph()
    ph.paragraph_format.space_before = Pt(6)
    ph.paragraph_format.space_after = Pt(2)
    add_run(ph, f"• {title}: ", bold=True, size=11, color=D700)
    add_run(ph, body, size=11, color=BODY)

# ── FAIR USE NOTE ──────────────────────────────────────────────────────────────
doc.add_paragraph()
fair = doc.add_paragraph()
fair.paragraph_format.space_before = Pt(8)
add_run(fair, '* Message fair-use limits at 3× the plan average. '
             'Enterprise pricing available on request — contact us for custom quotes.',
         italic=True, size=10, color=MUTED)

# ── SAVE ──────────────────────────────────────────────────────────────────────
out = "/Users/aliciapang/Documents/GitHub/Sequor/workspaces/_template/Sequor_Pricing_Strategy.docx"
doc.save(out)
print(f"Saved: {out}")
