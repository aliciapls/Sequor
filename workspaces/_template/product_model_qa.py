"""Generate Product Model Summary + Technical Q&A Word document."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Calibri"
D500 = RGBColor(0x29, 0x91, 0xAA)
D700 = RGBColor(0x17, 0x5A, 0x73)
D800 = RGBColor(0x14, 0x47, 0x5D)
BODY = RGBColor(0x4A, 0x5E, 0x6E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
MUTED = RGBColor(0x8A, 0x9E, 0xB0)
LIGHT = RGBColor(0xF4, 0xFB, 0xFD)

doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)


def shade_para(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def rule(doc, color='AAE5E8', space_before=4, space_after=8):
    hr = doc.add_paragraph()
    hr.paragraph_format.space_before = Pt(space_before)
    hr.paragraph_format.space_after = Pt(space_after)
    pPr = hr._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def run(para, text, bold=False, italic=False, size=11, color=BODY):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = FONT
    return r


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run(p, text, bold=True, size=16, color=D800)


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run(p, text, bold=True, size=13, color=D700)


def body(doc, text, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run(p, text, size=11, color=BODY)


# ─────────────────────────────────────────────────────────────────────────────
# PART 1 — PRODUCT MODEL
# ─────────────────────────────────────────────────────────────────────────────

# Top bar
bar = doc.add_paragraph()
bar.paragraph_format.space_before = Pt(0)
bar.paragraph_format.space_after = Pt(0)
shade_para(bar, '2991AA')
run(bar, '  ', size=6)

# Title
t1 = doc.add_paragraph()
t1.paragraph_format.space_before = Pt(18)
t1.paragraph_format.space_after = Pt(2)
run(t1, 'Product Model', bold=True, size=30, color=D800)

t2 = doc.add_paragraph()
t2.paragraph_format.space_after = Pt(2)
run(t2, 'Sequor — AI Coverage Layer for Customer Communication', italic=True, size=12, color=MUTED)

rule(doc, space_before=4, space_after=12)

body(doc,
     "Sequor is an AI-powered customer communication platform that operates as the "
     "always-on coverage layer between a business and its customers. It reads incoming "
     "WhatsApp and email messages, classifies them by intent, answers routine queries "
     "autonomously using a RAG pipeline grounded in the business's own documents, "
     "escalates what it cannot resolve, and learns from every human correction to "
     "improve over time.")

# ── Core Architecture ───────────────────────────────────────────────────────
h2(doc, "Core Architecture: Human-in-the-Loop with Confidence Badges")

body(doc,
     "Every AI-generated response carries a confidence badge. This is the foundational "
     "design choice — not full autonomy, not full manual review. The system operates on "
     "three confidence bands:")

bands = [
    ("High Confidence (>90%)", D500,
     "Auto-send. The AI resolves the query from the Document Hub and sends the response directly. "
     "No human involvement. The contact receives a reply within minutes."),
    ("Medium Confidence (60–90%)", RGBColor(0x54, 0xA8, 0xC0),
     "Human review. The AI drafts the response and flags it for the contact's review. "
     "The contact approves, edits, or overrides before it goes out. "
     "This is where the learning loop captures corrections."),
    ("Low Confidence (<60%)", RGBColor(0xE8, 0x8A, 0x00),
     "Escalation queue. The message is routed to the backup contact with full thread context. "
     "PII is redacted before context is delivered. SLA timer starts."),
]

for band_name, color, desc in bands:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    shade_para(p, 'F4FBFD')
    r1 = p.add_run(f"  {band_name}: ")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = color
    r1.font.name = FONT
    r2 = p.add_run(desc)
    r2.font.size = Pt(11)
    r2.font.color.rgb = BODY
    r2.font.name = FONT

# ── AI Pipeline ────────────────────────────────────────────────────────────
h2(doc, "AI Pipeline — 6-Step Flow")

steps = [
    ("1. Inbound",          "WhatsApp or email arrives via webhook. Normalized into a single unified contact + thread. Stored against the tenant's database."),
    ("2. Classify",         "A Kaizen LLM classifier labels the message by intent: Routine / Escalation / VIP. A confidence score (0–1) is attached. This score determines what happens next."),
    ("3. Retrieve",          "If Routine, the RAG pipeline queries the business's Document Hub. Top-k chunks are retrieved with citation. RAG is grounded exclusively in the business's own documents — no general web data."),
    ("4. Draft",            "A second LLM call generates the response from the retrieved chunks. The response is marked [DRAFT]. If confidence > 90%, it auto-sends. If 60–90%, it waits for human approval. If < 60%, it escalates."),
    ("5. Send or Hold",     "High confidence → auto-send via the original channel. Medium confidence → notification to the contact for approval. Low confidence → escalation queue with full thread context."),
    ("6. Learn",           "When a human edits, approves, or overrides an AI draft, the correction is logged to the Learning Loop. This feedback is used to retune classification thresholds and retrieval quality over time."),
]

for step_name, step_desc in steps:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run(p, f"{step_name}: ", bold=True, size=11, color=D500)
    run(p, step_desc, size=11, color=BODY)

# ── Channels ────────────────────────────────────────────────────────────────
h2(doc, "Channels")

channels = [
    ("WhatsApp",
     "Company-owned WhatsApp Business API account. Messages arrive via Meta webhook → Sequor portal. "
     "Template messages are pre-approved by Meta for OOO windows. The 24-hour WhatsApp session window "
     "is managed by acknowledging within the window and escalating what cannot be resolved."),
    ("Email",
     "SendGrid Inbound Parse webhook. Forward or BCC a Sequor address → webhook → portal. "
     "Full thread context is preserved. PDPA-compliant subject and body handling."),
    ("API / Webhooks",
     "REST API for custom integrations. Outbound webhooks to CRMs or helpdesks. "
     "Zapier / Make.com connectors for no-code integrations."),
]

for ch_name, ch_desc in channels:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    run(p, f"{ch_name}: ", bold=True, size=11, color=D700)
    run(p, ch_desc, size=11, color=BODY)

# ── Learning Loop ────────────────────────────────────────────────────────────
h2(doc, "Learning Loop — Compounding Moat")

body(doc,
     "The learning loop is the core compounding mechanism. Every human correction "
     "— an edit to an AI draft, an override, an escalation outcome — is logged as feedback. "
     "This feedback is used to:")

loop_points = [
    "Retune classification thresholds per tenant (some businesses use different jargon than others)",
    "Improve retrieval ranking in the RAG pipeline (which document chunks actually resolved queries)",
    "Adjust confidence calibration over time (if corrections cluster at 70–80% confidence, thresholds shift)",
]
for pt in loop_points:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run(p, f"  • {pt}", size=11, color=BODY)

body(doc,
     "The more businesses use Sequor, the more routing decisions it accumulates. "
     "A competitor starting from zero cannot replicate this — they can replicate the UI, "
     "but not 18 months of real routing feedback data.", space_after=8)

# ── PDPA Compliance ────────────────────────────────────────────────────────
h2(doc, "PDPA Compliance — Built In, Not Bolted On")

pillars = [
    ("PII Classification",    "Email addresses, phone numbers, and names are classified at ingestion. PII is redacted from RAG retrieval and classification records before storage."),
    ("Data Minimization",    "Only operationally necessary data is stored. Historical archives are pseudonymized."),
    ("Right to Erasure",   "Full deletion workflow. PII is purged from records and audit logs when a contact exercises their erasure right."),
    ("Audit Trail",         "Every action is logged: tenant, timestamp, actor, operation, outcome. Immutable. Retrievable on demand."),
    ("Tenant Isolation",     "Role-based access. Operators see only their own contacts. No cross-tenant data access."),
    ("Consent Management",  "Contact consent is flagged by channel and notice version. Retrievable on demand."),
]

for pill_name, pill_desc in pillars:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    run(p, f"{pill_name}: ", bold=True, size=11, color=D700)
    run(p, pill_desc, size=11, color=BODY)

body(doc,
     "SOC 2 Type II and ISO 27001 certification targeted for Q3 2026.", space_after=8)

# ── Multi-Tenancy ─────────────────────────────────────────────────────────
h2(doc, "Multi-Tenant Architecture")

body(doc,
     "Each business operates in an isolated tenant context. "
     "Documents, contacts, message threads, escalation rules, and audit logs are "
     "scoped to the tenant. The learning loop operates per-tenant initially, "
     "with cross-tenant anonymized aggregation used to improve default thresholds "
     "over time (opt-in).")

body(doc,
     "This architecture ensures that switching costs accumulate at the tenant level: "
     "a business that leaves Sequor loses its escalation history, its learned thresholds, "
     "its Document Hub, and its coverage record. These are not exportable in a usable form — "
     "they are Sequor's institutional memory.", space_after=16)


# ─────────────────────────────────────────────────────────────────────────────
# PAGE BREAK
# ─────────────────────────────────────────────────────────────────────────────
doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# PART 2 — TECHNICAL Q&A
# ─────────────────────────────────────────────────────────────────────────────

bar2 = doc.add_paragraph()
bar2.paragraph_format.space_before = Pt(0)
bar2.paragraph_format.space_after = Pt(0)
shade_para(bar2, '14475D')
run(bar2, '  ', size=6)

t3 = doc.add_paragraph()
t3.paragraph_format.space_before = Pt(18)
t3.paragraph_format.space_after = Pt(2)
run(t3, 'Technical Q&A', bold=True, size=30, color=D800)
t4 = doc.add_paragraph()
t4.paragraph_format.space_after = Pt(2)
run(t4, 'Questions you will be asked — and how to answer them', italic=True, size=12, color=MUTED)
rule(doc, '14475D', space_before=4, space_after=14)

qas = [
    (
        "What LLM do you use?",
        "We use Ollama for local inference, which gives us model portability "
        "and keeps data within the tenant's infrastructure. This means customer message "
        "content never leaves our infrastructure unless self-hosted by the customer. "
        "The specific model is configurable per deployment — we recommend Llama 3 or "
        "Mistral for the classification and drafting tasks. The RAG retrieval uses "
        "a separate embedding model (sentence-transformers) for chunk embeddings."
    ),
    (
        "How does the learning loop actually improve accuracy?",
        "Every human correction is logged with three pieces of data: the original message, "
        "the AI's draft, and the human's correction. This creates a labeled dataset per tenant. "
        "Periodically — not in real-time — we use this to retune the retrieval ranking "
        "(which chunks led to useful answers) and recalibrate the confidence thresholds "
        "(were corrections clustering in a specific band?). Cross-tenant anonymized patterns "
        "feed into the base model's default thresholds over time."
    ),
    (
        "What happens if the AI gives a wrong answer?",
        "The confidence badge architecture is the safety net. For high-confidence responses "
        "(>90%), the system auto-sends — but the operator sees all auto-sent responses in the audit log "
        "and can recall or correct them. For medium-confidence (60–90%), a human must approve "
        "before it goes out. For low-confidence (<60%), it never auto-sends — it escalates. "
        "There is no path where a low-confidence response goes to the customer unsupervised."
    ),
    (
        "How is PII handled in the RAG pipeline?",
        "PII is redacted at ingestion — before it enters classification or retrieval. "
        "Email addresses, phone numbers, NRIC numbers, and credit card patterns are detected "
        "and replaced with [REDACTED] tokens before storage. The RAG pipeline retrieves from "
        "redacted text only. The original unredacted message is stored in the AuditEntry "
        "table, which is immutable and access-controlled separately. "
        "This means a data subject erasure request can be honored — the PII can be purged "
        "from RAG/classification records while the audit trail remains intact."
    ),
    (
        "How does WhatsApp integration work without violating Meta's 24-hour window?",
        "The 24-hour window is managed proactively. For routine queries, the AI responds "
        "within minutes — well within the window. For queries that escalate, we send a "
        "template message (pre-approved by Meta) within the window that acknowledges receipt "
        "and sets expectations. The actual resolution comes via email or when the human "
        "becomes available. This is a designed tradeoff — not a gap."
    ),
    (
        "How does multi-tenancy work at the database level?",
        "Each tenant has an isolated schema (separate database schemas, not just row-level filtering). "
        "Documents, contacts, messages, and audit logs are all tenant-scoped. "
        "Cache keys include tenant_id. Audit rows persist tenant_id as an indexed column. "
        "A tenant_id is required for every read/write operation — the application layer enforces this, "
        "not just the query layer."
    ),
    (
        "What is the escalation routing logic?",
        "Escalation is triggered by: (1) confidence < 60%, (2) keywords matching a high-friction intent "
        "(refund threats, cancellation), (3) VIP contact flag, (4) SLA breach risk, or "
        "(5) a manual operator flag. When escalated, the message is assigned to the configured "
        "backup contact with: full thread context, retrieved RAG excerpts, SLA countdown, "
        "and PII redacted from the context delivered. The operator resolves from the portal. "
        "The outcome is logged and fed back to the learning loop."
    ),
    (
        "What infrastructure does this run on?",
        "The portal and API run on FastAPI (Python). The database is PostgreSQL via DataFlow. "
        "WhatsApp webhook processing is handled asynchronously. "
        "Email uses SendGrid Inbound Parse. "
        "RAG vector storage uses a local embedding model + PostgreSQL vector extension (pgvector). "
        "The current demo runs on localhost with ngrok exposing webhooks. "
        "Production deployment is containerized (Docker) with standard cloud hosting (AWS/GCP)."
    ),
    (
        "How do you prevent hallucinations from RAG?",
        "Three safeguards: (1) Confidence scoring — if retrieved chunks are low-quality or "
        "irrelevant, the drafting LLM's confidence reflects that and triggers escalation. "
        "(2) Citation requirement — every RAG-grounded answer must cite the source chunk. "
        "If no chunk is retrieved, the system defaults to escalation rather than drafting. "
        "(3) No general web data — RAG is grounded only in the business's own documents, "
        "which are controlled and auditable. No internet retrieval."
    ),
    (
        "What does SOC 2 Type II certification involve?",
        "SOC 2 Type II tests that our security controls operate effectively over a period "
        "(typically 6–12 months of audit observation). It covers: access controls, "
        "encryption in transit and at rest, incident response, change management, "
        "and monitoring. We are targeting Q3 2026. ISO 27001 (information security "
        "management) is a parallel track. This is the standard enterprise acquirers "
        "and regulated industry buyers will ask for."
    ),
    (
        "Why does the document cleanup service exist?",
        "Because RAG quality depends on document quality. Most SMEs don't have clean, "
        "structured documents — they have WhatsApp chat exports, informal notes, "
        "inconsistent spreadsheets. The cleanup service (a one-time $300–500 engagement) "
        "reviews, structures, and indexes the business's documents into RAG-ready format. "
        "Without this, the RAG pipeline returns irrelevant chunks and confidence drops. "
        "With it, the pipeline reliably answers routine queries. "
        "This is also a switching cost — once structured, the documents are Sequor's institutional memory."
    ),
    (
        "How does the learning loop avoid compounding errors?",
        "Each correction is logged with a confidence band, not just the final answer. "
        "If corrections cluster at a specific confidence level, that band's threshold "
        "is recalibrated. Corrections at high confidence are treated differently from "
        "corrections at low confidence. Additionally, the learning loop is per-tenant "
        "initially — a correction from one tenant doesn't shift the base model for "
        "another tenant. Cross-tenant aggregation is anonymized and used only for "
        "improving default thresholds."
    ),
]

for i, (question, answer) in enumerate(qas):
    # Q
    q_para = doc.add_paragraph()
    q_para.paragraph_format.space_before = Pt(10)
    q_para.paragraph_format.space_after = Pt(2)
    shade_para(q_para, 'D4EAF4')
    run(q_para, f"  Q{i+1}. {question}", bold=True, size=11, color=D800)

    # A
    a_para = doc.add_paragraph()
    a_para.paragraph_format.space_before = Pt(2)
    a_para.paragraph_format.space_after = Pt(6)
    run(a_para, f"  A: {answer}", size=11, color=BODY)


# ── SAVE ──────────────────────────────────────────────────────────────────────
out = "/Users/aliciapang/Documents/GitHub/Sequor/workspaces/_template/Sequor_Product_Model_and_Technical_QA.docx"
doc.save(out)
print(f"Saved: {out}")
