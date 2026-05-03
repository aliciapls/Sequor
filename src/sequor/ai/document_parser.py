"""Document parsing for various file formats.

Supports PDF, DOCX, XLSX, CSV, TXT, and image-based documents requiring OCR.
"""

import io
from abc import ABC, abstractmethod
from dataclasses import dataclass
from typing import Any

import structlog

logger = structlog.get_logger()


@dataclass
class ParsedDocument:
    """Result of document parsing."""

    text: str
    metadata: dict[str, Any]
    pages_failed: int = 0
    pages_total: int = 0


class DocumentParser(ABC):
    """Base class for document parsers."""

    @abstractmethod
    async def parse(self, content: bytes, filename: str) -> ParsedDocument:
        """Parse document content.

        Args:
            content: Raw file bytes
            filename: Original filename

        Returns:
            Parsed document with text and metadata
        """
        ...


class PDFParser(DocumentParser):
    """Parser for PDF documents."""

    async def parse(self, content: bytes, filename: str) -> ParsedDocument:
        """Extract text from PDF using pdfminer."""
        try:
            from pypdf import PdfReader
        except ImportError:
            logger.warning("pypdf.not_installed", filename=filename)
            return ParsedDocument(
                text="",
                metadata={"error": "pypdf not installed", "filename": filename},
            )

        text_parts = []
        pages_total = 0
        pages_failed = 0

        try:
            reader = PdfReader(io.BytesIO(content))
            pages_total = len(reader.pages)

            for page_num, page in enumerate(reader.pages):
                try:
                    text = page.extract_text()
                    if text:
                        text_parts.append(f"[Page {page_num + 1}]\n{text}")
                    else:
                        pages_failed += 1
                        logger.debug(
                            "pdf.page.empty",
                            filename=filename,
                            page=page_num + 1,
                        )
                except Exception as e:
                    pages_failed += 1
                    logger.warning(
                        "pdf.page.failed",
                        filename=filename,
                        page=page_num + 1,
                        error=str(e),
                    )

        except Exception as e:
            logger.error("pdf.parse.failed", filename=filename, error=str(e))
            return ParsedDocument(
                text="",
                metadata={"error": str(e), "filename": filename},
                pages_failed=1,
                pages_total=1,
            )

        text = "\n\n".join(text_parts)
        logger.info(
            "pdf.parse.ok",
            filename=filename,
            pages_total=pages_total,
            pages_failed=pages_failed,
            text_length=len(text),
        )

        return ParsedDocument(
            text=text,
            metadata={
                "filename": filename,
                "pages_total": pages_total,
                "pages_failed": pages_failed,
                "parser": "pypdf",
            },
            pages_failed=pages_failed,
            pages_total=pages_total,
        )


class DOCXParser(DocumentParser):
    """Parser for DOCX documents."""

    async def parse(self, content: bytes, filename: str) -> ParsedDocument:
        """Extract text from DOCX using python-docx."""
        try:
            import docx
        except ImportError:
            logger.warning("python-docx.not_installed", filename=filename)
            return ParsedDocument(
                text="",
                metadata={"error": "python-docx not installed", "filename": filename},
            )

        try:
            doc = docx.Document(io.BytesIO(content))
            paragraphs = []

            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            text = "\n".join(paragraphs)

            logger.info(
                "docx.parse.ok",
                filename=filename,
                paragraph_count=len(paragraphs),
                text_length=len(text),
            )

            return ParsedDocument(
                text=text,
                metadata={
                    "filename": filename,
                    "paragraph_count": len(paragraphs),
                    "parser": "python-docx",
                },
            )

        except Exception as e:
            logger.error("docx.parse.failed", filename=filename, error=str(e))
            return ParsedDocument(
                text="",
                metadata={"error": str(e), "filename": filename},
            )


class XLSXParser(DocumentParser):
    """Parser for XLSX spreadsheets."""

    async def parse(self, content: bytes, filename: str) -> ParsedDocument:
        """Extract text from XLSX using openpyxl."""
        try:
            import openpyxl
        except ImportError:
            logger.warning("openpyxl.not_installed", filename=filename)
            return ParsedDocument(
                text="",
                metadata={"error": "openpyxl not installed", "filename": filename},
            )

        try:
            wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
            rows_data = []
            sheet_count = 0

            for sheet_name in wb.sheetnames:
                sheet_count += 1
                ws = wb[sheet_name]
                rows_data.append(f"[Sheet: {sheet_name}]")

                for row_num, row in enumerate(ws.iter_rows(values_only=True), 1):
                    if any(cell is not None for cell in row):
                        row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                        rows_data.append(f"Row {row_num}: {row_text}")

            text = "\n".join(rows_data)

            logger.info(
                "xlsx.parse.ok",
                filename=filename,
                sheet_count=sheet_count,
                text_length=len(text),
            )

            return ParsedDocument(
                text=text,
                metadata={
                    "filename": filename,
                    "sheet_count": sheet_count,
                    "parser": "openpyxl",
                },
            )

        except Exception as e:
            logger.error("xlsx.parse.failed", filename=filename, error=str(e))
            return ParsedDocument(
                text="",
                metadata={"error": str(e), "filename": filename},
            )


class CSVParser(DocumentParser):
    """Parser for CSV files."""

    async def parse(self, content: bytes, filename: str) -> ParsedDocument:
        """Extract text from CSV using built-in csv module."""
        import csv

        try:
            text_lines = []
            row_count = 0

            decoded_content = content.decode("utf-8", errors="replace")
            reader = csv.DictReader(io.StringIO(decoded_content))

            headers = reader.fieldnames or []
            text_lines.append("Columns: " + ", ".join(headers))

            for row in reader:
                row_count += 1
                row_text = " | ".join(f"{k}: {v}" for k, v in row.items() if v)
                text_lines.append(f"Row: {row_text}")

            text = "\n".join(text_lines)

            logger.info(
                "csv.parse.ok",
                filename=filename,
                row_count=row_count,
                text_length=len(text),
            )

            return ParsedDocument(
                text=text,
                metadata={
                    "filename": filename,
                    "row_count": row_count,
                    "column_count": len(headers),
                    "parser": "csv",
                },
            )

        except Exception as e:
            logger.error("csv.parse.failed", filename=filename, error=str(e))
            return ParsedDocument(
                text="",
                metadata={"error": str(e), "filename": filename},
            )


class TXTParser(DocumentParser):
    """Parser for plain text files."""

    async def parse(self, content: bytes, filename: str) -> ParsedDocument:
        """Extract text from plain text."""
        try:
            text = content.decode("utf-8", errors="replace")

            logger.info(
                "txt.parse.ok",
                filename=filename,
                text_length=len(text),
            )

            return ParsedDocument(
                text=text,
                metadata={
                    "filename": filename,
                    "parser": "txt",
                },
            )

        except Exception as e:
            logger.error("txt.parse.failed", filename=filename, error=str(e))
            return ParsedDocument(
                text="",
                metadata={"error": str(e), "filename": filename},
            )


class OCRParser(DocumentParser):
    """Parser for image-based documents using OCR."""

    async def parse(self, content: bytes, filename: str) -> ParsedDocument:
        """Extract text from images using pytesseract.

        Requires Tesseract OCR to be installed on the system.
        """
        try:
            import pytesseract
            from PIL import Image
        except ImportError as e:
            logger.warning(
                "ocr.dependencies.not_installed",
                filename=filename,
                error=str(e),
            )
            return ParsedDocument(
                text="",
                metadata={
                    "error": "OCR dependencies not installed (pytesseract/PIL)",
                    "filename": filename,
                },
            )

        try:
            image = Image.open(io.BytesIO(content))
            text = pytesseract.image_to_string(image)

            logger.info(
                "ocr.parse.ok",
                filename=filename,
                text_length=len(text),
            )

            return ParsedDocument(
                text=text,
                metadata={
                    "filename": filename,
                    "parser": "tesseract",
                    "image_format": image.format,
                },
            )

        except Exception as e:
            logger.error("ocr.parse.failed", filename=filename, error=str(e))
            return ParsedDocument(
                text="",
                metadata={"error": str(e), "filename": filename},
            )


def get_parser_for_file(filename: str) -> DocumentParser:
    """Get the appropriate parser for a file based on its extension.

    Args:
        filename: The filename to get a parser for

    Returns:
        Appropriate DocumentParser instance
    """
    ext = filename.lower().split(".")[-1] if "." in filename else ""

    parsers = {
        "pdf": PDFParser(),
        "docx": DOCXParser(),
        "xlsx": XLSXParser(),
        "xls": XLSXParser(),
        "csv": CSVParser(),
        "txt": TXTParser(),
        "png": OCRParser(),
        "jpg": OCRParser(),
        "jpeg": OCRParser(),
    }

    parser = parsers.get(ext)
    if parser is None:
        logger.warning("parser.unknown_extension", filename=filename, extension=ext)
        return TXTParser()

    return parser
